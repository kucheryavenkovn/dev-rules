# FILE: src/renderer.py
# VERSION: 1.3.0
# START_MODULE_CONTRACT
#   PURPOSE: Рендеринг HTML-контента в элементы docx: параграфы, таблицы, списки, код, изображения
#   SCOPE: render_html_to_doc, render_paragraph, render_table, render_list, render_image
#   DEPENDS: M-CONFIG, M-TYPES
#   LINKS: M-RENDERER
#   ROLE: RUNTIME
#   MAP_MODE: EXPORTS
# END_MODULE_CONTRACT
#
# START_MODULE_MAP
#   render_html_to_doc - главная функция рендеринга HTML → docx (bookmark_map, chapter_rel_path)
#   render_paragraph - рендеринг параграфа с inline-стилями и внутренними ссылками
#   render_table - рендеринг HTML таблицы в Word таблицу
#   render_list - рендеринг списка с вложенностью (явная нумерация, перезапуск)
#   render_image - встраивание изображения из docs/ в docx
#   _resolve_internal_link - разрешить .md href → GRACE bookmark name
#   _add_hyperlink - создать w:hyperlink с url (external) или anchor (internal)
# END_MODULE_MAP

import logging
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString

from src.types import ChapterInfo

from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def _add_hyperlink(paragraph, text, url=None, anchor=None):
    """Добавить кликабельную гиперссылку в параграф docx."""
    hyperlink = OxmlElement("w:hyperlink")
    if anchor:
        hyperlink.set(qn("w:anchor"), anchor)
    elif url:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# START_BLOCK_RENDER_HTML
def render_html_to_doc(doc, html_content, depth_offset=0, _recursion_depth=0, img_base_dir=None, bookmark_map=None, chapter_rel_path=None):
    """Рендерить HTML-контент в параграфы docx."""
    if _recursion_depth > 5:
        return
    logger.debug("[Renderer][render_html_to_doc][BLOCK_RENDER_HTML] depth_offset=%d recursion=%d len=%d", depth_offset, _recursion_depth, len(html_content))
    soup = BeautifulSoup(html_content, "html.parser")

    for element in soup.children:
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                doc.add_paragraph(text)
            continue

        tag = element.name

        if tag in ("h1", "h2", "h3", "h4"):
            level = int(tag[1]) + depth_offset
            level = min(level, 4)
            text = element.get_text(strip=True)
            if text:
                doc.add_heading(text, level=level)

        elif tag == "p":
            render_paragraph(doc, element, img_base_dir=img_base_dir, bookmark_map=bookmark_map, chapter_rel_path=chapter_rel_path)

        elif tag == "pre":
            code_el = element.find("code")
            text = code_el.get_text() if code_el else element.get_text()
            for line in text.splitlines():
                p = doc.add_paragraph(line)
                p.style = doc.styles["Code"]
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.left_indent = Cm(0.5)

        elif tag == "table":
            render_table(doc, element)

        elif tag in ("ul", "ol"):
            render_list(doc, element, ordered=(tag == "ol"))

        elif tag == "blockquote":
            text = element.get_text(strip=True)
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Cm(1.0)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        elif tag == "hr":
            pass

        elif tag in ("div", "section"):
            render_html_to_doc(doc, str(element), depth_offset, _recursion_depth + 1, img_base_dir=img_base_dir, bookmark_map=bookmark_map, chapter_rel_path=chapter_rel_path)

        elif tag == "img":
            src = element.get("src", "")
            render_image(doc, src, img_base_dir)

        elif tag in ("svg", "mermaid"):
            pass

        else:
            text = element.get_text(strip=True)
            if text:
                doc.add_paragraph(text)
# END_BLOCK_RENDER_HTML


# START_BLOCK_RENDER_PARAGRAPH
def _resolve_internal_link(href, bookmark_map, chapter_rel_path=None):
    """Разрешить внутреннюю ссылку в bookmark name."""
    if not bookmark_map:
        return None
    clean = href.lstrip("./")
    if "#" in clean:
        clean = clean.split("#")[0]
    if clean in bookmark_map:
        return bookmark_map[clean]
    if chapter_rel_path:
        from pathlib import PurePosixPath
        base = PurePosixPath(chapter_rel_path).parent
        resolved = str(base / clean).replace("\\", "/")
        if resolved in bookmark_map:
            return bookmark_map[resolved]
    return None


def render_paragraph(doc, element, img_base_dir=None, bookmark_map=None, chapter_rel_path=None):
    """Рендерить HTML параграф с inline-стилями."""
    p = doc.add_paragraph()
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip():
                p.add_run(text)
        elif child.name in ("strong", "b"):
            r = p.add_run(child.get_text())
            r.bold = True
        elif child.name in ("em", "i"):
            r = p.add_run(child.get_text())
            r.italic = True
        elif child.name == "code":
            r = p.add_run(child.get_text())
            r.font.name = "Courier New"
            r.font.size = Pt(8)
        elif child.name == "a":
            text = child.get_text()
            href = child.get("href", "")
            if text.strip():
                if href.startswith("http://") or href.startswith("https://"):
                    _add_hyperlink(p, text, url=href)
                else:
                    anchor = _resolve_internal_link(href, bookmark_map, chapter_rel_path)
                    if anchor:
                        _add_hyperlink(p, text, anchor=anchor)
                        logger.debug("[Renderer][render_paragraph][BLOCK_RENDER_LINK] mode=anchor target=%s href=%s", anchor, href)
                    else:
                        r = p.add_run(text)
                        r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                        r.underline = True
                        logger.debug("[Renderer][render_paragraph][BLOCK_RENDER_LINK] mode=styled href=%s", href)
        elif child.name == "br":
            pass
        elif child.name == "img":
            src = child.get("src", "")
            render_image(doc, src, img_base_dir)
        else:
            text = child.get_text()
            if text.strip():
                p.add_run(text)
    return p
# END_BLOCK_RENDER_PARAGRAPH


# START_BLOCK_RENDER_LIST
def render_list(doc, element, ordered=False, level=0):
    """Рендерить HTML список."""
    items = element.find_all("li", recursive=False)
    for idx, li in enumerate(items):
        code_blocks = li.find_all("pre", recursive=False)
        inline_code = li.find_all("code", recursive=False)
        sublists = li.find_all(["ul", "ol"], recursive=False)
        has_code_content = code_blocks or inline_code

        if has_code_content:
            for child in li.children:
                if isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        indent = Cm(0.5 * (level + 1))
                        prefix = f"{idx + 1}. " if ordered else ""
                        p = doc.add_paragraph(f"{prefix}{text}")
                        p.paragraph_format.left_indent = indent
                        p.style = doc.styles["Normal"]
                elif child.name == "pre":
                    code_el = child.find("code")
                    text = code_el.get_text() if code_el else child.get_text()
                    for line in text.splitlines():
                        p = doc.add_paragraph(line)
                        p.style = doc.styles["Code"]
                        pf = p.paragraph_format
                        pf.space_before = Pt(0)
                        pf.space_after = Pt(0)
                        pf.left_indent = Cm(0.5 * (level + 2))
                elif child.name in ("ul", "ol"):
                    pass
                elif child.name == "code":
                    p = doc.add_paragraph(child.get_text())
                    p.style = doc.styles["Code"]
                    pf = p.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                    pf.left_indent = Cm(0.5 * (level + 2))
                elif child.name == "p":
                    text_parts = []
                    runs_data = []
                    for sub in child.children:
                        if isinstance(sub, NavigableString):
                            t = str(sub)
                            if t.strip():
                                runs_data.append({"text": t, "code": False})
                        elif sub.name == "code":
                            runs_data.append({"text": sub.get_text(), "code": True})
                        else:
                            t = sub.get_text()
                            if t.strip():
                                runs_data.append({"text": t, "code": False})
                    indent = Cm(0.5 * (level + 1))
                    prefix = f"{idx + 1}. " if ordered else ""
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = indent
                    p.style = doc.styles["Normal"]
                    if runs_data:
                        p.add_run(prefix)
                        for rd in runs_data:
                            r = p.add_run(rd["text"])
                            if rd["code"]:
                                r.font.name = "Courier New"
                                r.font.size = Pt(8)
                else:
                    text = child.get_text(strip=True)
                    if text:
                        indent = Cm(0.5 * (level + 1))
                        prefix = f"{idx + 1}. " if ordered else ""
                        p = doc.add_paragraph(f"{prefix}{text}")
                        p.paragraph_format.left_indent = indent
                        p.style = doc.styles["Normal"]
        else:
            text = li.get_text(strip=True)
            indent = Cm(0.5 * (level + 1))
            if ordered:
                p = doc.add_paragraph(f"{idx + 1}. {text}")
            else:
                p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = indent
            p.style = doc.styles["Normal"]

        for sub in sublists:
            render_list(doc, sub, ordered=(sub.name == "ol"), level=level + 1)
# END_BLOCK_RENDER_LIST


def _render_cell_content(cell_element):
    """Извлечь runs из HTML-ячейки с учётом inline <code>."""
    runs = []
    for child in cell_element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip():
                runs.append({"text": text, "code": False})
        elif child.name == "code":
            runs.append({"text": child.get_text(), "code": True})
        else:
            text = child.get_text()
            if text.strip():
                runs.append({"text": text, "code": False})
    return runs


# START_BLOCK_RENDER_TABLE
def render_table(doc, element):
    """Рендерить HTML таблицу."""
    rows = element.find_all("tr")
    if not rows:
        logger.debug("[Renderer][render_table][BLOCK_RENDER_TABLE] skipped=no_rows")
        return

    first_row = rows[0]
    cols = first_row.find_all(["td", "th"])
    num_cols = len(cols)
    if num_cols == 0:
        logger.debug("[Renderer][render_table][BLOCK_RENDER_TABLE] skipped=no_cols")
        return

    table = doc.add_table(rows=min(len(rows), 100), cols=num_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        is_header = row.find("th") is not None
        for j, cell in enumerate(cells):
            if j >= num_cols:
                continue
            table_cell = table.rows[i].cells[j]
            table_cell.text = ""
            p = table_cell.paragraphs[0]
            has_code = cell.find("code") is not None
            if has_code:
                runs_data = _render_cell_content(cell)
                for rd in runs_data:
                    r = p.add_run(rd["text"])
                    r.font.size = Pt(10)
                    if is_header or cell.name == "th":
                        r.bold = True
                    if rd["code"]:
                        r.font.name = "Courier New"
                        r.font.size = Pt(8)
            else:
                text = cell.get_text(strip=True)
                p.text = text
                for run in p.runs:
                    run.font.size = Pt(10)
                    if is_header or cell.name == "th":
                        run.bold = True

    logger.info("[Renderer][render_table][BLOCK_RENDER_TABLE] rows=%d cols=%d", len(rows), num_cols)
    doc.add_paragraph()
# END_BLOCK_RENDER_TABLE


# START_BLOCK_RENDER_IMAGE
def render_image(doc, src, img_base_dir=None):
    """Встроить изображение в docx или вставить плейсхолдер."""
    if not src:
        return

    image_path = None
    if img_base_dir is not None:
        candidate = Path(img_base_dir) / src
        if candidate.exists():
            image_path = candidate

    if image_path is not None:
        try:
            doc.add_picture(str(image_path), width=Inches(5.5))
            logger.info("[Renderer][render_image][BLOCK_RENDER_IMAGE] embedded=%s", src)
        except Exception as e:
            logger.warning("[Renderer][render_image][BLOCK_RENDER_IMAGE] failed=%s error=%s", src, e)
            doc.add_paragraph(f"[Image: {src}]")
    else:
        logger.warning("[Renderer][render_image][BLOCK_RENDER_IMAGE] not_found=%s", src)
        doc.add_paragraph(f"[Image: {src}]")
# END_BLOCK_RENDER_IMAGE

# START_CHANGE_SUMMARY
#   LAST_CHANGE: v1.3.0 — render_html_to_doc accepts chapter_rel_path, passes to render_paragraph for internal cross-reference resolution via bookmark_map
# END_CHANGE_SUMMARY
