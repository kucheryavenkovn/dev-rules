# FILE: src/main.py
# VERSION: 2.0.0
# START_MODULE_CONTRACT
#   PURPOSE: Точка входа генератора Word-документа с GRACE-разметкой
#   SCOPE: main — сборка документа: титульная → оглавление → главы → GRACE инъекция → валидация
#   DEPENDS: M-CONFIG, M-RENDERER, M-SIDEBAR, M-GRACE, M-VALIDATOR, M-DISCOVERY, M-GRAPHSYNC, M-TYPES
#   LINKS: M-DOCXBUILDER
#   ROLE: ENTRY_POINT
#   MAP_MODE: EXPORTS
# END_MODULE_CONTRACT
#
# START_MODULE_MAP
#   main - точка входа: синхронизация графа → генерация docx → GRACE инъекция → валидация
# END_MODULE_MAP

import sys
import io
import logging
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.config import (
    setup_styles, classify_module_type, DOC_NAME, DOC_VERSION,
    OUTPUT_PATH, GRACE_VERSION, derive_module_id,
)
from src.parser import read_md, md_to_html
from src.sidebar_order import build_chapter_order
from src.types import ChapterInfo
from src.renderer import render_html_to_doc
from src.grace_injector import inject_grace_parts, inject_bookmark_start, inject_bookmark_end
from src.validator import validate_grace_docx
from src.graph_sync import sync_graph_from_fs

logging.basicConfig(level=logging.INFO, format="%(message)s")
logger = logging.getLogger(__name__)


def _make_run_rpr(font_size_pt, bold=False, color_val='1F3A5F', underline=False):
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size_pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(font_size_pt * 2)))
    rPr.append(szCs)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement('w:b'))
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color_val)
    rPr.append(c)
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    return rPr


def add_toc_hyperlink(paragraph, bookmark_name, text, font_size_pt, bold=False):
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)
    new_run = OxmlElement('w:r')
    new_run.append(_make_run_rpr(font_size_pt, bold=bold, color_val='1F3A5F', underline=True))
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


def add_pageref_field(paragraph, bookmark_name, font_size_pt):
    r_begin = OxmlElement('w:r')
    fc_begin = OxmlElement('w:fldChar')
    fc_begin.set(qn('w:fldCharType'), 'begin')
    r_begin.append(fc_begin)
    paragraph._element.append(r_begin)

    r_instr = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' PAGEREF {bookmark_name} \\h '
    r_instr.append(instrText)
    paragraph._element.append(r_instr)

    r_sep = OxmlElement('w:r')
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    r_sep.append(fc_sep)
    paragraph._element.append(r_sep)

    r_text = OxmlElement('w:r')
    r_text.append(_make_run_rpr(font_size_pt, color_val='1F3A5F'))
    t = OxmlElement('w:t')
    t.text = '0'
    r_text.append(t)
    paragraph._element.append(r_text)

    r_end = OxmlElement('w:r')
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    r_end.append(fc_end)
    paragraph._element.append(r_end)


def add_toc_tab_stop(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:leader'), 'dot')
    tab.set(qn('w:pos'), '9072')
    tabs.append(tab)
    pPr.append(tabs)


def add_tab_run(paragraph):
    r = OxmlElement('w:r')
    r.append(OxmlElement('w:tab'))
    paragraph._element.append(r)


def main():
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    today_str = date.today().isoformat()

    print("=== Генерация Word-документа с GRACE-разметкой ===\n")

    # START_BLOCK_SYNC_GRAPH
    print("[1/6] Синхронизация knowledge-graph.xml с файловой системой...")
    flat = sync_graph_from_fs()
    print(f"       Обнаружено {len(flat)} глав в docs/")
    # END_BLOCK_SYNC_GRAPH

    # START_BLOCK_BUILD_DOC
    print("[2/6] Сборка документа Word...")
    doc = Document()
    setup_styles(doc)

    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("СТАНДАРТЫ РАЗРАБОТКИ")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle_p.add_run("Гид по внутренним процессам и стандартам\nпроектной команды")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()

    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ver_p.add_run(f"Версия {DOC_VERSION}")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = date_p.add_run(today_str)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()
    doc.add_paragraph()

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = org_p.add_run("yellow-hammer / dev-rules")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()
    # END_BLOCK_BUILD_DOC

    # START_BLOCK_TOC
    print("[3/6] Построение оглавления...")
    chapters = build_chapter_order(use_graph=True)
    if not chapters:
        print("       Fallback: chapters пустые, используем discovery напрямую")
        chapters = build_chapter_order(use_graph=False)

    existing_ids = set()
    resolved = []
    for ch in chapters:
        mid = ch.module_id or derive_module_id(ch.heading, existing_ids)
        existing_ids.add(mid)
        resolved.append(ch.with_module_id(mid))
    chapters = resolved

    toc_heading = doc.add_heading("Оглавление", level=1)

    for ch in chapters:
        toc_p = doc.add_paragraph()
        toc_p.paragraph_format.left_indent = Cm(ch.depth * 0.75)

        bookmark_name = f"GRACE_{ch.module_id}"
        font_size = 12 if ch.depth == 0 else 11
        is_bold = ch.depth == 0

        add_toc_tab_stop(toc_p)
        add_toc_hyperlink(toc_p, bookmark_name, ch.heading, font_size, bold=is_bold)
        add_tab_run(toc_p)
        add_pageref_field(toc_p, bookmark_name, font_size)

    doc.add_page_break()
    # END_BLOCK_TOC

    # START_BLOCK_CHAPTERS
    print(f"[4/6] Генерация {len(chapters)} глав...")
    para_counter = 0
    table_counter = 0
    h1_counter = 0
    h2_counter = 0
    bookmark_id = 100
    modules_info = []

    for ch in chapters:
        mid = ch.module_id
        depth = ch.depth
        heading_level = min(depth + 1, 4)

        heading_para = doc.add_heading(ch.heading, level=heading_level)

        inject_bookmark_start(heading_para, bookmark_id, mid)

        para_start = para_counter
        current_bookmark_id = bookmark_id
        bookmark_id += 1
        para_counter += 1

        if heading_level == 1:
            h1_counter += 1
        elif heading_level == 2:
            h2_counter += 1

        mod_type = "NARRATIVE"
        mod_elements = []

        if ch.source:
            title, md_text = read_md(ch.source)
            if md_text:
                html = md_to_html(md_text)
                mod_type = classify_module_type(html)

                html_tables = html.lower().count("<table>")
                table_counter += html_tables
                for t_idx in range(html_tables):
                    mod_elements.append({
                        "type": "TABLE-DATA",
                        "para-index": str(para_counter + t_idx),
                        "columns": "0",
                        "rows": "0",
                    })

                render_html_to_doc(doc, html, depth_offset=depth)
                para_counter += len(html.split("<p>")) + len(html.split("<h")) + len(html.split("<li>"))

        else:
            mod_type = "NAVIGATION"

        para_end = para_counter - 1

        inject_bookmark_end(doc, current_bookmark_id)

        modules_info.append({
            "id": mid,
            "heading": ch.heading,
            "type": mod_type,
            "para_start": para_start,
            "para_end": para_end,
            "elements": mod_elements,
            "subsections": [],
            "parent": ch.parent_heading,
            "parent_heading": ch.parent_heading,
        })
    # END_BLOCK_CHAPTERS

    # START_BLOCK_SAVE_AND_INJECT
    print("[5/6] Сохранение и GRACE-инъекция...")
    doc.save(str(OUTPUT_PATH))

    grace_output = OUTPUT_PATH.parent / f"GRACE_{OUTPUT_PATH.name}"
    inject_grace_parts(
        base_docx_path=OUTPUT_PATH,
        grace_docx_path=grace_output,
        modules_info=modules_info,
        total_paras=para_counter,
        total_tables=table_counter,
        total_h1=h1_counter,
        total_h2=h2_counter,
        today_str=today_str,
    )
    # END_BLOCK_SAVE_AND_INJECT

    # START_BLOCK_VALIDATE
    print("[6/6] Валидация...")
    result = validate_grace_docx(grace_output)
    status = "OK" if result["valid"] else "FAILED"
    print(f"       Валидация: {status}")
    if not result["valid"]:
        for err in result["errors"]:
            print(f"       ERROR: {err}")
    # END_BLOCK_VALIDATE

    # START_BLOCK_REPORT
    print(f"""
=============================================
GRACE-DOCX Bootstrap Complete  [v3]
=============================================
Document:        {DOC_NAME}
Version:         {DOC_VERSION}
Modules:         {len(modules_info)}
XML parts:       5 injected
Bookmarks:       {len(modules_info)} pairs
Graph synced:    docs/knowledge-graph.xml → grace-graph.xml
---------------------------------------------""")

    for mod in modules_info:
        print(f"  {mod['id']:12s}  {mod['type']:12s}  {mod['heading']}")

    print(f"""---------------------------------------------
Output:
  Base:    {OUTPUT_PATH}
  GRACE:   {grace_output}
  Status:  {status}
============================================""")
    # END_BLOCK_REPORT


if __name__ == "__main__":
    main()
