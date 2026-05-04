# FILE: tests/unit/test_grace_injector.py
# VERSION: 1.0.0
# START_MODULE_CONTRACT
#   PURPOSE: Юнит-тесты модуля M-GRACE (src/grace_injector.py)
#   SCOPE: make_grace_manifest, make_grace_instructions, make_grace_contracts, make_grace_verification, inject_grace_parts, inject_bookmark_start/end
#   DEPENDS: M-GRACE
#   LINKS: V-M-GRACE
#   ROLE: TEST
#   MAP_MODE: LOCALS
# END_MODULE_CONTRACT

import zipfile

from lxml import etree

from docx import Document

from src.grace_injector import (
    make_grace_manifest,
    make_grace_instructions,
    make_grace_contracts,
    make_grace_verification,
    inject_grace_parts,
    inject_bookmark_start,
    inject_bookmark_end,
    GRACE_PART_NAMES,
)
from src.config import setup_styles


def _parse_xml(xml_string):
    return etree.fromstring(xml_string.encode("utf-8"))


class TestMakeGraceManifest:
    def test_well_formed_xml(self):
        xml = make_grace_manifest([], "2025-01-01")
        root = _parse_xml(xml)
        assert root.tag == "GraceManifest"

    def test_contains_document_name(self):
        xml = make_grace_manifest([], "2025-01-01")
        root = _parse_xml(xml)
        name_el = root.find("document-name")
        assert name_el is not None
        assert name_el.text

    def test_contains_protocol(self):
        xml = make_grace_manifest([], "2025-01-01")
        root = _parse_xml(xml)
        protocol = root.find("Protocol")
        assert protocol is not None
        steps = [el for el in protocol if el.tag.startswith("step-")]
        assert len(steps) == 9

    def test_contains_parts(self):
        xml = make_grace_manifest([], "2025-01-01")
        root = _parse_xml(xml)
        parts = root.find("Parts")
        assert parts is not None
        part_els = [el for el in parts if el.tag.startswith("part-")]
        assert len(part_els) == 5


class TestMakeGraceInstructions:
    def test_well_formed_xml(self):
        xml = make_grace_instructions()
        root = _parse_xml(xml)
        assert root.tag == "GraceInstructions"

    def test_has_core_principles(self):
        xml = make_grace_instructions()
        root = _parse_xml(xml)
        principles = root.find("CorePrinciples")
        assert principles is not None
        items = [el for el in principles if el.tag.startswith("principle-")]
        assert len(items) >= 6

    def test_has_edit_rules(self):
        xml = make_grace_instructions()
        root = _parse_xml(xml)
        rules = root.find("EditRules")
        assert rules is not None


class TestMakeGraceContracts:
    def test_well_formed_xml(self):
        modules = [
            {"id": "M-TEST", "heading": "Test Module", "type": "NARRATIVE"},
        ]
        xml = make_grace_contracts(modules)
        root = _parse_xml(xml)
        assert root.tag == "GraceContracts"

    def test_has_type_contracts(self):
        xml = make_grace_contracts([])
        root = _parse_xml(xml)
        tc = root.find("TypeContracts")
        assert tc is not None
        assert tc.find("C-NARRATIVE") is not None
        assert tc.find("C-TABLE-DATA") is not None
        assert tc.find("C-MIXED") is not None

    def test_module_contract_generated(self):
        modules = [
            {"id": "M-TEST", "heading": "Test", "type": "NARRATIVE"},
        ]
        xml = make_grace_contracts(modules)
        root = _parse_xml(xml)
        mc = root.find("ModuleContracts")
        assert mc is not None
        assert mc.find("C-M-TEST") is not None

    def test_empty_modules_no_crash(self):
        xml = make_grace_contracts([])
        root = _parse_xml(xml)
        assert root.tag == "GraceContracts"


class TestMakeGraceVerification:
    def test_well_formed_xml(self):
        xml = make_grace_verification(10)
        root = _parse_xml(xml)
        assert root.tag == "GraceVerification"

    def test_has_structural_invariants(self):
        xml = make_grace_verification(10)
        root = _parse_xml(xml)
        si = root.find("StructuralInvariants")
        assert si is not None

    def test_bookmark_balance_invariant(self):
        xml = make_grace_verification(76)
        assert 'bookmark-balance' in xml
        assert '76' in xml

    def test_has_validation_protocol(self):
        xml = make_grace_verification(10)
        root = _parse_xml(xml)
        vp = root.find("ValidationProtocol")
        assert vp is not None


class TestInjectBookmarks:
    def test_bookmark_start_injected(self):
        doc = Document()
        p = doc.add_paragraph("Test")
        inject_bookmark_start(p, 100, "M-TEST")

        xml = p._element.xml
        assert "bookmarkStart" in xml
        assert "GRACE_M-TEST" in xml

    def test_bookmark_end_injected(self):
        doc = Document()
        doc.add_paragraph("Test")
        inject_bookmark_end(doc, 100)

        body_xml = doc.element.body.xml
        assert "bookmarkEnd" in body_xml


class TestInjectGraceParts:
    def test_creates_grace_docx(self, tmp_path):
        doc = Document()
        setup_styles(doc)
        doc.add_paragraph("Test content")
        base = tmp_path / "base.docx"
        grace = tmp_path / "GRACE_base.docx"
        doc.save(str(base))

        modules = [
            {"id": "M-TEST", "heading": "Test", "type": "NARRATIVE",
             "para_start": 0, "para_end": 0, "elements": [], "subsections": [],
             "parent": None, "parent_heading": None},
        ]

        inject_grace_parts(base, grace, modules, 1, 0, 1, 0, "2025-01-01")

        assert grace.exists()

    def test_all_xml_parts_present(self, tmp_path):
        doc = Document()
        setup_styles(doc)
        doc.add_paragraph("Test")
        base = tmp_path / "base.docx"
        grace = tmp_path / "GRACE_base.docx"
        doc.save(str(base))

        modules = [
            {"id": "M-TEST", "heading": "Test", "type": "NARRATIVE",
             "para_start": 0, "para_end": 0, "elements": [], "subsections": [],
             "parent": None, "parent_heading": None},
        ]
        inject_grace_parts(base, grace, modules, 1, 0, 1, 0, "2025-01-01")

        with zipfile.ZipFile(str(grace), "r") as zf:
            names = zf.namelist()
        for name in GRACE_PART_NAMES:
            assert f"word/{name}" in names, f"Missing: word/{name}"

    def test_content_types_updated(self, tmp_path):
        doc = Document()
        setup_styles(doc)
        doc.add_paragraph("Test")
        base = tmp_path / "base.docx"
        grace = tmp_path / "GRACE_base.docx"
        doc.save(str(base))

        modules = [
            {"id": "M-TEST", "heading": "Test", "type": "NARRATIVE",
             "para_start": 0, "para_end": 0, "elements": [], "subsections": [],
             "parent": None, "parent_heading": None},
        ]
        inject_grace_parts(base, grace, modules, 1, 0, 1, 0, "2025-01-01")

        with zipfile.ZipFile(str(grace), "r") as zf:
            ct = zf.read("[Content_Types].xml").decode("utf-8")
        for name in GRACE_PART_NAMES:
            assert name in ct, f"Missing Content_Types entry for {name}"

    def test_rels_updated(self, tmp_path):
        doc = Document()
        setup_styles(doc)
        doc.add_paragraph("Test")
        base = tmp_path / "base.docx"
        grace = tmp_path / "GRACE_base.docx"
        doc.save(str(base))

        modules = [
            {"id": "M-TEST", "heading": "Test", "type": "NARRATIVE",
             "para_start": 0, "para_end": 0, "elements": [], "subsections": [],
             "parent": None, "parent_heading": None},
        ]
        inject_grace_parts(base, grace, modules, 1, 0, 1, 0, "2025-01-01")

        with zipfile.ZipFile(str(grace), "r") as zf:
            rels = zf.read("word/_rels/document.xml.rels").decode("utf-8")
        for i in range(1, 6):
            assert f"rIdGrace{i}" in rels

    def test_all_grace_xml_well_formed(self, tmp_path):
        doc = Document()
        setup_styles(doc)
        doc.add_paragraph("Test")
        base = tmp_path / "base.docx"
        grace = tmp_path / "GRACE_base.docx"
        doc.save(str(base))

        modules = [
            {"id": "M-TEST", "heading": "Test", "type": "NARRATIVE",
             "para_start": 0, "para_end": 0, "elements": [], "subsections": [],
             "parent": None, "parent_heading": None},
        ]
        inject_grace_parts(base, grace, modules, 1, 0, 1, 0, "2025-01-01")

        with zipfile.ZipFile(str(grace), "r") as zf:
            for name in GRACE_PART_NAMES:
                content = zf.read(f"word/{name}")
                etree.fromstring(content)

    def test_empty_modules_info(self, tmp_path):
        doc = Document()
        setup_styles(doc)
        doc.add_paragraph("Test")
        base = tmp_path / "base.docx"
        grace = tmp_path / "GRACE_base.docx"
        doc.save(str(base))

        inject_grace_parts(base, grace, [], 0, 0, 0, 0, "2025-01-01")
        assert grace.exists()
