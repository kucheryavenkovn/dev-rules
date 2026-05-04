# FILE: tests/unit/test_renderer.py
# VERSION: 1.1.0
# START_MODULE_CONTRACT
#   PURPOSE: Юнит-тесты модуля M-RENDERER (src/renderer.py)
#   SCOPE: render_html_to_doc, render_paragraph, render_table, render_list, render_image
#   DEPENDS: M-RENDERER
#   LINKS: V-M-RENDERER
#   ROLE: TEST
#   MAP_MODE: LOCALS
# END_MODULE_CONTRACT

import struct
import zlib

import logging
from pathlib import Path

import pytest
from docx import Document

from src.config import setup_styles
from src.renderer import (
    render_html_to_doc,
    render_paragraph,
    render_table,
    render_list,
)


def _make_doc():
    doc = Document()
    setup_styles(doc)
    return doc


def _write_valid_png(path, width=4, height=4):
    def _chunk(chunk_type, data):
        c = chunk_type + data
        crc = struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)
        return struct.pack(">I", len(data)) + c + crc

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr_data = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    ihdr = _chunk(b"IHDR", ihdr_data)
    raw = b""
    for y in range(height):
        raw += b"\x00" + b"\xff\x00" * width
    idat = _chunk(b"IDAT", zlib.compress(raw))
    iend = _chunk(b"IEND", b"")
    path.write_bytes(sig + ihdr + idat + iend)


class TestRenderTable:
    def test_basic_table_rows_and_cols(self):
        doc = _make_doc()
        html = "<table><tr><th>A</th><th>B</th></tr><tr><td>1</td><td>2</td></tr></table>"
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        table_el = soup.find("table")
        render_table(doc, table_el)

        assert len(doc.tables) == 1
        t = doc.tables[0]
        assert len(t.rows) == 2
        assert len(t.columns) == 2

    def test_table_cell_content(self):
        doc = _make_doc()
        html = "<table><tr><td>Hello</td><td>World</td></tr></table>"
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        render_table(doc, soup.find("table"))

        t = doc.tables[0]
        assert t.rows[0].cells[0].text == "Hello"
        assert t.rows[0].cells[1].text == "World"

    def test_empty_table_no_crash(self):
        doc = _make_doc()
        html = "<table></table>"
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        render_table(doc, soup.find("table"))
        assert len(doc.tables) == 0

    def test_table_no_rows_no_crash(self):
        doc = _make_doc()
        html = "<table><tr></tr></table>"
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        render_table(doc, soup.find("table"))
        assert len(doc.tables) == 0

    def test_header_row_bold(self):
        doc = _make_doc()
        html = "<table><tr><th>Header</th></tr><tr><td>Data</td></tr></table>"
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        render_table(doc, soup.find("table"))

        t = doc.tables[0]
        header_runs = t.rows[0].cells[0].paragraphs[0].runs
        assert any(r.bold for r in header_runs)


class TestRenderCode:
    def test_code_block_uses_code_style(self):
        doc = _make_doc()
        html = "<pre><code>print('hello')</code></pre>"
        render_html_to_doc(doc, html)

        code_paras = [p for p in doc.paragraphs if p.style and p.style.name == "Code"]
        assert len(code_paras) >= 1

    def test_code_block_multiline(self):
        doc = _make_doc()
        html = "<pre><code>line1\nline2\nline3</code></pre>"
        render_html_to_doc(doc, html)

        code_paras = [p for p in doc.paragraphs if p.style and p.style.name == "Code"]
        assert len(code_paras) == 3


class TestRenderParagraph:
    def test_bold_inline(self):
        doc = _make_doc()
        html = "<p>Hello <strong>world</strong></p>"
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        p_el = soup.find("p")
        render_paragraph(doc, p_el)

        last_para = doc.paragraphs[-1]
        bold_runs = [r for r in last_para.runs if r.bold]
        assert len(bold_runs) >= 1

    def test_italic_inline(self):
        doc = _make_doc()
        html = "<p>Hello <em>world</em></p>"
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        p_el = soup.find("p")
        render_paragraph(doc, p_el)

        last_para = doc.paragraphs[-1]
        italic_runs = [r for r in last_para.runs if r.italic]
        assert len(italic_runs) >= 1

    def test_code_inline(self):
        doc = _make_doc()
        html = "<p>Use <code>pip install</code> to install</p>"
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        p_el = soup.find("p")
        render_paragraph(doc, p_el)

        last_para = doc.paragraphs[-1]
        code_runs = [r for r in last_para.runs if r.font.name == "Courier New"]
        assert len(code_runs) >= 1

    def test_link_inline(self):
        doc = _make_doc()
        html = '<p><a href="https://example.com">click here</a></p>'
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        p_el = soup.find("p")
        render_paragraph(doc, p_el)

        last_para = doc.paragraphs[-1]
        hyperlinks = last_para._p.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
        )
        assert len(hyperlinks) >= 1
        assert hyperlinks[0].text == "click here"


class TestRenderList:
    def test_unordered_list(self):
        doc = _make_doc()
        html = "<ul><li>Item 1</li><li>Item 2</li></ul>"
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        ul_el = soup.find("ul")
        render_list(doc, ul_el, ordered=False)

        list_items = [p for p in doc.paragraphs if p.text.strip() in ("Item 1", "Item 2")]
        assert len(list_items) >= 2

    def test_ordered_list(self):
        doc = _make_doc()
        html = "<ol><li>First</li><li>Second</li></ol>"
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        ol_el = soup.find("ol")
        render_list(doc, ol_el, ordered=True)

        list_items = [p for p in doc.paragraphs if "First" in p.text or "Second" in p.text]
        assert len(list_items) >= 2
        assert list_items[0].text.strip().startswith("1.")

    def test_nested_list(self):
        doc = _make_doc()
        html = "<ul><li>Outer<ul><li>Inner</li></ul></li></ul>"
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        ul_el = soup.find("ul")
        render_list(doc, ul_el, ordered=False)

        texts = [p.text.strip() for p in doc.paragraphs]
        assert any("Inner" in t for t in texts), f"Inner not found in {texts}"


class TestRenderHtmlToDoc:
    def test_headings_with_depth_offset(self):
        doc = _make_doc()
        html = "<h1>Title</h1><h2>Subtitle</h2>"
        render_html_to_doc(doc, html, depth_offset=1)

        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) == 2

    def test_blockquote(self):
        doc = _make_doc()
        html = "<blockquote>A quote</blockquote>"
        render_html_to_doc(doc, html)

        quote_paras = [p for p in doc.paragraphs if "A quote" in p.text]
        assert len(quote_paras) >= 1

    def test_recursion_limit(self):
        doc = _make_doc()
        html = '<div><div><div><div><div><div><p>Deep</p></div></div></div></div></div></div>'
        render_html_to_doc(doc, html, _recursion_depth=5)

        deep_paras = [p for p in doc.paragraphs if "Deep" in p.text]
        assert len(deep_paras) == 0

    def test_hr_skipped(self):
        doc = _make_doc()
        html = "<p>Before</p><hr><p>After</p>"
        render_html_to_doc(doc, html)

        texts = [p.text for p in doc.paragraphs]
        assert "Before" in texts
        assert "After" in texts

    def test_img_embedded_from_base_dir(self, tmp_path):
        img_file = tmp_path / "test.png"
        _write_valid_png(img_file)
        doc = _make_doc()
        html = '<img src="test.png"/>'
        render_html_to_doc(doc, html, img_base_dir=tmp_path)
        inline_shapes = doc.inline_shapes
        assert len(inline_shapes) >= 1

    def test_img_missing_produces_placeholder(self):
        doc = _make_doc()
        html = '<img src="nonexistent.png"/>'
        render_html_to_doc(doc, html, img_base_dir=Path("/tmp"))
        placeholder_paras = [p for p in doc.paragraphs if "[Image:" in p.text]
        assert len(placeholder_paras) >= 1

    def test_img_subdirectory_resolved(self, tmp_path):
        sub = tmp_path / "img"
        sub.mkdir()
        img_file = sub / "photo.png"
        _write_valid_png(img_file)
        doc = _make_doc()
        html = '<img src="img/photo.png"/>'
        render_html_to_doc(doc, html, img_base_dir=tmp_path)
        inline_shapes = doc.inline_shapes
        assert len(inline_shapes) >= 1

    def test_full_html_document(self):
        doc = _make_doc()
        html = """
        <h1>Title</h1>
        <p>Paragraph with <strong>bold</strong> and <em>italic</em></p>
        <table><tr><th>H</th></tr><tr><td>D</td></tr></table>
        <ul><li>Item</li></ul>
        <pre><code>x = 1</code></pre>
        """
        render_html_to_doc(doc, html)

        assert len(doc.tables) == 1
        assert any(p.style.name == "Code" for p in doc.paragraphs)


class TestInternalCrossReference:
    def test_internal_link_to_bookmark(self):
        doc = _make_doc()
        html = '<p><a href="./metadata/roles.md">Роли</a></p>'
        bookmark_map = {"metadata/roles.md": "GRACE_M-ROLES"}
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        p_el = soup.find("p")
        render_paragraph(doc, p_el, bookmark_map=bookmark_map)

        last_para = doc.paragraphs[-1]
        hyperlinks = last_para._p.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
        )
        assert len(hyperlinks) == 1
        assert hyperlinks[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor") == "GRACE_M-ROLES"
        assert hyperlinks[0].text == "Роли"

    def test_internal_link_not_found_fallback_styled(self):
        doc = _make_doc()
        html = '<p><a href="./nonexistent.md">Unknown</a></p>'
        bookmark_map = {"metadata/roles.md": "GRACE_M-ROLES"}
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        p_el = soup.find("p")
        render_paragraph(doc, p_el, bookmark_map=bookmark_map)

        last_para = doc.paragraphs[-1]
        hyperlinks = last_para._p.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
        )
        assert len(hyperlinks) == 0
        assert any(r.text == "Unknown" for r in last_para.runs)

    def test_external_link_unchanged(self):
        doc = _make_doc()
        html = '<p><a href="https://example.com">external</a></p>'
        bookmark_map = {"metadata/roles.md": "GRACE_M-ROLES"}
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        p_el = soup.find("p")
        render_paragraph(doc, p_el, bookmark_map=bookmark_map)

        last_para = doc.paragraphs[-1]
        hyperlinks = last_para._p.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
        )
        assert len(hyperlinks) == 1
        has_rid = any("id" in attr for attr in hyperlinks[0].attrib)
        assert has_rid
        assert hyperlinks[0].text == "external"

    def test_internal_link_relative_from_subdir(self):
        doc = _make_doc()
        html = '<p><a href="../begin.md">Начало</a></p>'
        bookmark_map = {"begin.md": "GRACE_M-BEGIN"}
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        p_el = soup.find("p")
        render_paragraph(doc, p_el, bookmark_map=bookmark_map, chapter_rel_path="cicd/README.md")

        last_para = doc.paragraphs[-1]
        hyperlinks = last_para._p.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
        )
        assert len(hyperlinks) == 1
        assert hyperlinks[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor") == "GRACE_M-BEGIN"

    def test_internal_link_via_render_html_to_doc(self):
        doc = _make_doc()
        html = '<p>См. <a href="./metadata/roles.md">Роли</a></p>'
        bookmark_map = {"metadata/roles.md": "GRACE_M-ROLES"}
        render_html_to_doc(doc, html, bookmark_map=bookmark_map)

        last_para = doc.paragraphs[-1]
        hyperlinks = last_para._p.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
        )
        assert len(hyperlinks) == 1
        assert hyperlinks[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor") == "GRACE_M-ROLES"

    def test_no_bookmark_map_no_crash(self):
        doc = _make_doc()
        html = '<p><a href="./metadata/roles.md">Роли</a></p>'
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, "html.parser")
        p_el = soup.find("p")
        render_paragraph(doc, p_el)

        last_para = doc.paragraphs[-1]
        assert any("Роли" in (r.text or "") for r in last_para.runs)


class TestListNumberingRestart:
    def test_two_ol_restart_at_one(self):
        doc = _make_doc()
        html = "<ol><li>A</li><li>B</li></ol><p>text</p><ol><li>C</li><li>D</li></ol>"
        render_html_to_doc(doc, html)

        ol_paras = [
            p for p in doc.paragraphs
            if p.text.strip() in ("A", "B", "C", "D")
            or p.text.strip().startswith(("1.", "2."))
        ]
        texts = [p.text.strip() for p in ol_paras]
        first_numbers = [t for t in texts if t.startswith("1.")]
        assert len(first_numbers) >= 2, f"Expected at least two '1.' prefixes, got: {texts}"

    def test_ol_uses_explicit_number_not_list_number_style(self):
        doc = _make_doc()
        html = "<ol><li>Item</li></ol>"
        render_html_to_doc(doc, html)

        ol_paras = [
            p for p in doc.paragraphs
            if "Item" in p.text
        ]
        assert len(ol_paras) >= 1
        p = ol_paras[0]
        assert p.style.name != "List Number", f"Should NOT use List Number style, got: {p.style.name}"
        assert p.text.strip().startswith("1."), f"Should start with explicit '1.', got: {p.text}"
