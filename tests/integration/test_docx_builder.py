# FILE: tests/integration/test_docx_builder.py
# VERSION: 1.0.0
# START_MODULE_CONTRACT
#   PURPOSE: Интеграционные тесты сквозного цикла DF-001: Markdown → docx → GRACE инъекция → валидация
#   SCOPE: Full generation pipeline — M-DOCXBUILDER
#   DEPENDS: M-DOCXBUILDER, M-RENDERER, M-STRUCTURE, M-GRACE, M-VALIDATOR
#   LINKS: V-M-DOCXBUILDER, VF-001
#   ROLE: TEST
#   MAP_MODE: LOCALS
# END_MODULE_CONTRACT

import zipfile

from lxml import etree

from docx import Document

from src.config import setup_styles, classify_module_type, derive_module_id
from src.parser import read_md, md_to_html
from src.structure import collect_chapters
from src.renderer import render_html_to_doc
from src.grace_injector import inject_grace_parts, inject_bookmark_start, inject_bookmark_end, GRACE_PART_NAMES
from src.validator import validate_grace_docx
from src.graph_sync import sync_graph_from_fs


def _build_docx(tmp_path, chapter_limit=None):
    base_docx = tmp_path / "base.docx"
    grace_docx = tmp_path / "GRACE_base.docx"

    sync_graph_from_fs()
    chapters = collect_chapters(use_graph=True)
    if not chapters:
        chapters = collect_chapters(use_graph=False)

    if chapter_limit:
        chapters = chapters[:chapter_limit]

    doc = Document()
    setup_styles(doc)

    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = 1
    r = title_p.add_run("TEST DOCUMENT")
    r.font.size = 20
    r.font.bold = True

    doc.add_page_break()

    bookmark_id = 100
    modules_info = []
    para_counter = 0
    table_counter = 0
    h1_counter = 0
    h2_counter = 0

    for ch in chapters:
        mid = ch.get("module_id") or derive_module_id(ch["heading"], {m["id"] for m in modules_info})
        depth = ch.get("depth", 0)
        heading_level = min(depth + 1, 4)

        heading_para = doc.add_heading(ch["heading"], level=heading_level)
        inject_bookmark_start(heading_para, bookmark_id, mid)

        ch["bookmark_start_id"] = bookmark_id
        ch["para_start"] = para_counter
        bookmark_id += 1
        para_counter += 1

        if heading_level == 1:
            h1_counter += 1
        elif heading_level == 2:
            h2_counter += 1

        mod_type = "NARRATIVE"
        mod_elements = []

        if ch.get("source"):
            title, md_text = read_md(ch["source"])
            if md_text:
                html = md_to_html(md_text)
                mod_type = classify_module_type(html)
                html_tables = html.lower().count("<table>")
                table_counter += html_tables
                render_html_to_doc(doc, html, depth_offset=depth)
                para_counter += 10
        else:
            mod_type = "NAVIGATION"

        ch["type"] = mod_type
        ch["para_end"] = para_counter - 1

        inject_bookmark_end(doc, ch["bookmark_start_id"])

        modules_info.append({
            "id": mid,
            "heading": ch["heading"],
            "type": mod_type,
            "para_start": ch["para_start"],
            "para_end": ch["para_end"],
            "elements": mod_elements,
            "subsections": [],
            "parent": ch.get("parent_heading"),
            "parent_heading": ch.get("parent_heading"),
        })

    doc.save(str(base_docx))

    from datetime import date
    inject_grace_parts(
        base_docx_path=base_docx,
        grace_docx_path=grace_docx,
        modules_info=modules_info,
        total_paras=para_counter,
        total_tables=table_counter,
        total_h1=h1_counter,
        total_h2=h2_counter,
        today_str=date.today().isoformat(),
    )

    return base_docx, grace_docx, modules_info


class TestFullGenerationSmall:
    def test_creates_both_files(self, tmp_path):
        base, grace, modules = _build_docx(tmp_path, chapter_limit=5)
        assert base.exists()
        assert grace.exists()

    def test_bookmarks_balanced(self, tmp_path):
        base, grace, modules = _build_docx(tmp_path, chapter_limit=5)
        with zipfile.ZipFile(str(grace), "r") as zf:
            doc_xml = zf.read("word/document.xml").decode("utf-8")
        starts = doc_xml.count("<w:bookmarkStart")
        ends = doc_xml.count("<w:bookmarkEnd")
        assert starts == ends
        assert starts > 0

    def test_grace_bookmarks_present(self, tmp_path):
        base, grace, modules = _build_docx(tmp_path, chapter_limit=5)
        with zipfile.ZipFile(str(grace), "r") as zf:
            doc_xml = zf.read("word/document.xml").decode("utf-8")
        assert "GRACE_" in doc_xml

    def test_all_grace_parts_present(self, tmp_path):
        base, grace, modules = _build_docx(tmp_path, chapter_limit=5)
        with zipfile.ZipFile(str(grace), "r") as zf:
            names = zf.namelist()
        for name in GRACE_PART_NAMES:
            assert f"word/{name}" in names

    def test_all_grace_xml_well_formed(self, tmp_path):
        base, grace, modules = _build_docx(tmp_path, chapter_limit=5)
        with zipfile.ZipFile(str(grace), "r") as zf:
            for pname in GRACE_PART_NAMES:
                content = zf.read(f"word/{pname}")
                etree.fromstring(content)

    def test_content_types_complete(self, tmp_path):
        base, grace, modules = _build_docx(tmp_path, chapter_limit=5)
        with zipfile.ZipFile(str(grace), "r") as zf:
            ct = zf.read("[Content_Types].xml").decode("utf-8")
        for name in GRACE_PART_NAMES:
            assert name in ct

    def test_rels_complete(self, tmp_path):
        base, grace, modules = _build_docx(tmp_path, chapter_limit=5)
        with zipfile.ZipFile(str(grace), "r") as zf:
            rels = zf.read("word/_rels/document.xml.rels").decode("utf-8")
        for i in range(1, 6):
            assert f"rIdGrace{i}" in rels

    def test_validator_passes(self, tmp_path):
        base, grace, modules = _build_docx(tmp_path, chapter_limit=5)
        result = validate_grace_docx(grace)
        assert result["valid"], f"Validation failed: {result.get('errors', [])}"

    def test_module_count_matches_chapters(self, tmp_path):
        base, grace, modules = _build_docx(tmp_path, chapter_limit=5)
        assert len(modules) == 5

    def test_grace_graph_has_modules(self, tmp_path):
        base, grace, modules = _build_docx(tmp_path, chapter_limit=5)
        with zipfile.ZipFile(str(grace), "r") as zf:
            graph_xml = zf.read("word/grace-graph.xml")
        root = etree.fromstring(graph_xml)
        modules_el = root.find("Modules")
        assert modules_el is not None
        children = [el for el in modules_el if el.tag.startswith("M-")]
        assert len(children) == 5


class TestFullGenerationAll:
    def test_all_chapters(self, tmp_path):
        base, grace, modules = _build_docx(tmp_path)
        assert base.exists()
        assert grace.exists()
        assert len(modules) > 10

    def test_all_bookmarks_paired(self, tmp_path):
        base, grace, modules = _build_docx(tmp_path)
        with zipfile.ZipFile(str(grace), "r") as zf:
            doc_xml = zf.read("word/document.xml").decode("utf-8")
        starts = doc_xml.count("<w:bookmarkStart")
        ends = doc_xml.count("<w:bookmarkEnd")
        assert starts == ends

    def test_full_validation(self, tmp_path):
        base, grace, modules = _build_docx(tmp_path)
        result = validate_grace_docx(grace)
        assert result["valid"], f"Validation failed: {result.get('errors', [])}"
