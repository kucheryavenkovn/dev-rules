#!/usr/bin/env python3
"""
Генератор Word-документа «Стандарты разработки» с GRACE-разметкой.
Читает Markdown-файлы из docs/ и собирает .docx с семантическими закладками и XML-частями GRACE v3.
"""

import os
import re
import json
import copy
from datetime import date
from pathlib import Path
from collections import OrderedDict

import markdown
from bs4 import BeautifulSoup, NavigableString
from lxml import etree

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: F401 — используется ниже
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# ─── Конфигурация ───────────────────────────────────────────────────
DOCS_DIR = Path(r"D:\git\dev-rules\docs")
OUTPUT_PATH = Path(r"D:\git\dev-rules\Стандарты_разработки.docx")
DOC_NAME = "Стандарты разработки"
DOC_VERSION = "1.0"
GRACE_VERSION = "3.0.0"
TODAY = date.today().isoformat()

# ─── Утилиты ────────────────────────────────────────────────────────

def strip_frontmatter(text):
    """Удалить YAML frontmatter из markdown."""
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) >= 3:
            return parts[2].strip()
    return text


def get_frontmatter_title(text):
    """Извлечь заголовок из frontmatter (slug / title)."""
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) >= 3:
            fm = parts[1]
            for line in fm.splitlines():
                if line.strip().startswith("title:"):
                    return line.split(":", 1)[1].strip().strip('"').strip("'")
    return None


def read_md(rel_path):
    """Прочитать markdown файл и вернуть (title, clean_text)."""
    fpath = DOCS_DIR / rel_path
    if not fpath.exists():
        return None, None
    text = fpath.read_text(encoding="utf-8")
    title = get_frontmatter_title(text)
    clean = strip_frontmatter(text)
    if title is None:
        for line in clean.splitlines():
            if line.startswith("# "):
                title = line[2:].strip()
                break
    return title, clean


def md_to_html(md_text):
    """Конвертировать Markdown → HTML через библиотеку markdown."""
    return markdown.markdown(
        md_text,
        extensions=["tables", "fenced_code", "toc", "codehilite"],
        extension_defaults={"codehilite": {"guess_lang": False}},
    )


# ─── Определение структуры глав ─────────────────────────────────────
# Порядок и иерархия глав соответствует sidebars.js и _category_.json

CHAPTERS = OrderedDict([
    ("Введение", "intro.md"),
    ("Начало разработки", "begin.md"),
    ("Оформление кода", "layout.md"),
    ("Запросы", "request.md"),
    ("Управляемые формы", "forms.md"),
    ("Блокировка форм", "blocking_forms.md"),
    ("Ввод на основании", "input_based_on.md"),
    ("Расширения", "extensions.md"),
    ("Печатные формы", "printed_forms.md"),
    ("Префикс и комментарии", "prefix_comments.md"),
    ("Принципы эффективной разработки", OrderedDict([
        ("Обзор", "principles/README.md"),
        ("KISS — Делай проще", "principles/kiss.md"),
        ("DRY — Не повторяйся", "principles/dry.md"),
        ("YAGNI — Вам это не понадобится", "principles/yagni.md"),
        ("SOLID", "principles/solid.md"),
        ("Архитектура 1С-решений", "principles/architecture.md"),
    ])),
    ("Метаданные", OrderedDict([
        ("Общие модули", "metadata/common_modules.md"),
        ("Справочники", "metadata/catalogs.md"),
        ("Документы", "metadata/documents.md"),
        ("Регистры накопления", "metadata/accumulation_registers.md"),
        ("Регистры сведений", "metadata/information_registers.md"),
        ("Подсистемы", "metadata/subsystem.md"),
        ("Отчеты", "metadata/reports.md"),
        ("Роли", "metadata/roles.md"),
        ("Подписки на события", "metadata/event_subscriptions.md"),
        ("Регламентные задания", "metadata/scheduled_jobs.md"),
    ])),
    ("Система управления версиями", OrderedDict([
        ("Обзор", "version-control/README.md"),
        ("Хранилище 1С", "version-control/1c-storage/README.md"),
        ("Работа с Git", OrderedDict([
            ("Обзор", "version-control/git/README.md"),
            ("Основные команды Git", "version-control/git/commands.md"),
            ("Настройка SSH ключей", "version-control/git/ssh.md"),
            ("Git Flow", "version-control/git/gitflow.md"),
            ("Настройка исключений Git", "version-control/git/gitignore.md"),
            ("Настройка атрибутов Git", "version-control/git/gitattributes.md"),
            ("Git LFS", "version-control/git/lfs.md"),
            ("Подмодули Git", "version-control/git/submodules.md"),
        ])),
    ])),
    ("Среды разработки", OrderedDict([
        ("Обзор", "ide/README.md"),
        ("Phoenix BSL", "ide/phoenix-bsl.md"),
        ("Visual Studio Code", "ide/vscode.md"),
        ("1С:EDT", "ide/edt.md"),
    ])),
    ("DevOps", OrderedDict([
        ("Обзор", "cicd/README.md"),
        ("Code-review", "cicd/code-review/README.md"),
        ("Конвейеры CI/CD", OrderedDict([
            ("Обзор", "cicd/pipelines/README.md"),
            ("Профиль Jenkins", "cicd/pipelines/jenkins-pipeline-profile.md"),
            ("Профиль GitLab CI", "cicd/pipelines/gitlab-ci-pipeline-profile.md"),
        ])),
        ("Тестирование", OrderedDict([
            ("Обзор", "cicd/testing/README.md"),
            ("Стратегия тестирования", "cicd/testing/automation-strategy.md"),
            ("Инициализация тестовой ИБ", "cicd/testing/data-initialization.md"),
            ("Чек-лист ЗУП", "cicd/testing/hrm-checklist.md"),
        ])),
        ("Доставка и развертывание", "cicd/delivery/README.md"),
        ("SonarQube", "cicd/sonar/README.md"),
    ])),
    ("Интеграции", OrderedDict([
        ("Обзор", "integrations/README.md"),
        ("Брокеры сообщений", OrderedDict([
            ("Обзор", "integrations/messaging/README.md"),
            ("RabbitMQ", OrderedDict([
                ("Обзор", "integrations/messaging/rabbitmq/README.md"),
                ("Основные концепции", "integrations/messaging/rabbitmq/concepts.md"),
                ("Регламент именования", "integrations/messaging/rabbitmq/naming.md"),
                ("БИТ.Адаптер", "integrations/messaging/rabbitmq/bit-adapter/README.md"),
            ])),
            ("Apache Kafka", "integrations/messaging/kafka/README.md"),
        ])),
    ])),
    ("Руководства пользователя", OrderedDict([
        ("Обзор", "manuals/README.md"),
        ("Окружения", "manuals/environments.md"),
        ("Пользователи", "manuals/users.md"),
        ("Обновление конфигураций", "manuals/update_regulations.md"),
        ("Жизненный цикл задачи", "manuals/task-lifecycle.md"),
    ])),
    ("Глоссарий терминов", "glossary.md"),
])


# ─── Модульные ID для GRACE ─────────────────────────────────────────

def derive_module_id(heading, existing_ids):
    """Генерация Module ID из заголовка."""
    special = {
        "Введение": "M-INTRO",
        "Начало разработки": "M-BEGIN",
        "Оформление кода": "M-LAYOUT",
        "Запросы": "M-QUERY",
        "Управляемые формы": "M-FORMS",
        "Блокировка форм": "M-BLOCK",
        "Ввод на основании": "M-INPUT",
        "Расширения": "M-EXT",
        "Печатные формы": "M-PRINT",
        "Префикс и комментарии": "M-PREFIX",
        "Принципы эффективной разработки": "M-PRINC",
        "Метаданные": "M-META",
        "Система управления версиями": "M-VC",
        "Среды разработки": "M-IDE",
        "DevOps": "M-DEVOPS",
        "Интеграции": "M-INTEG",
        "Руководства пользователя": "M-MANUAL",
        "Глоссарий терминов": "M-GLOSS",
        "KISS — Делай проще": "M-KISS",
        "DRY — Не повторяйся": "M-DRY",
        "YAGNI — Вам это не понадобится": "M-YAGNI",
        "SOLID": "M-SOLID",
        "Архитектура 1С-решений": "M-ARCH",
        "Общие модули": "M-COMMON",
        "Справочники": "M-CATALOG",
        "Документы": "M-DOCS",
        "Регистры накопления": "M-ACCREG",
        "Регистры сведений": "M-INFREG",
        "Подсистемы": "M-SUBSYS",
        "Отчеты": "M-REPORTS",
        "Роли": "M-ROLES",
        "Подписки на события": "M-EVTSUB",
        "Регламентные задания": "M-SCHED",
        "Обзор": "M-OVERVIEW",
        "Хранилище 1С": "M-1CSTOR",
        "Работа с Git": "M-GIT",
        "Основные команды Git": "M-GITCMD",
        "Настройка SSH ключей": "M-SSH",
        "Git Flow": "M-GITFLOW",
        "Настройка исключений Git": "M-GITIGN",
        "Настройка атрибутов Git": "M-GITATTR",
        "Git LFS": "M-LFS",
        "Подмодули Git": "M-SUBMOD",
        "Phoenix BSL": "M-PHOENIX",
        "Visual Studio Code": "M-VSCODE",
        "1С:EDT": "M-EDT",
        "Code-review": "M-CR",
        "Конвейеры CI/CD": "M-PIPE",
        "Профиль Jenkins": "M-JENKINS",
        "Профиль GitLab CI": "M-GITLAB",
        "Тестирование": "M-TEST",
        "Стратегия тестирования": "M-TESTSTR",
        "Инициализация тестовой ИБ": "M-TESTDATA",
        "Чек-лист ЗУП": "M-HRMCHK",
        "Доставка и развертывание": "M-DELIVER",
        "SonarQube": "M-SONAR",
        "Брокеры сообщений": "M-MSG",
        "RabbitMQ": "M-RMQ",
        "Основные концепции": "M-RMQCON",
        "Регламент именования": "M-RMQNAM",
        "БИТ.Адаптер": "M-BITADPT",
        "Apache Kafka": "M-KAFKA",
        "Окружения": "M-ENV",
        "Пользователи": "M-USERS",
        "Обновление конфигураций": "M-UPDATE",
        "Жизненный цикл задачи": "M-TASK",
    }
    mid = special.get(heading)
    if mid:
        if mid not in existing_ids:
            return mid
        base = mid
        counter = 2
        while f"{base}-{counter}" in existing_ids:
            counter += 1
        return f"{base}-{counter}"
    clean = re.sub(r"[^A-Za-z0-9]", "", heading).upper()
    if len(clean) >= 3:
        candidate = "M-" + clean[:5]
    else:
        candidate = "M-SEC" + str(len(existing_ids) + 100)
    while candidate in existing_ids:
        candidate = candidate + str(len(existing_ids))
    return candidate


# ─── Определение типа модуля GRACE ──────────────────────────────────

def classify_module_type(html_content):
    """Определить тип модуля по содержимому."""
    has_tables = "<table>" in html_content.lower()
    has_code = "<code>" in html_content.lower() or "<pre>" in html_content.lower()
    has_prose = "<p>" in html_content.lower()
    if has_tables and has_code and has_prose:
        return "MIXED"
    if has_tables and has_code:
        return "MIXED"
    if has_tables:
        return "DATA"
    if has_code or has_prose:
        return "NARRATIVE"
    return "NARRATIVE"


# ─── Создание стилей документа ──────────────────────────────────────

def setup_styles(doc):
    """Настроить стили документа."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for i in range(1, 5):
        sname = f"Heading {i}"
        if sname in doc.styles:
            s = doc.styles[sname]
            s.font.name = "Calibri"
            if i == 1:
                s.font.size = Pt(22)
                s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
                s.font.bold = True
            elif i == 2:
                s.font.size = Pt(16)
                s.font.color.rgb = RGBColor(0x2C, 0x5F, 0x2D)
                s.font.bold = True
            elif i == 3:
                s.font.size = Pt(13)
                s.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
                s.font.bold = True
            elif i == 4:
                s.font.size = Pt(11)
                s.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                s.font.bold = True
            s.paragraph_format.space_before = Pt(12)
            s.paragraph_format.space_after = Pt(6)

    if "Code" not in [s.name for s in doc.styles]:
        cs = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        cs.font.name = "Consolas"
        cs.font.size = Pt(9)
        cs.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        cs.paragraph_format.space_before = Pt(2)
        cs.paragraph_format.space_after = Pt(2)
        cs.paragraph_format.left_indent = Cm(0.5)


# ─── Рендеринг HTML → docx ──────────────────────────────────────────

def render_html_to_doc(doc, html_content, depth_offset=0):
    """Рендерить HTML-контент в параграфы docx."""
    soup = BeautifulSoup(html_content, "html.parser")

    for element in soup.children:
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                doc.add_paragraph(text)
            continue

        tag = element.name

        if tag in ("h1", "h2", "h3", "h4"):
            level = int(tag[1]) + depth_offset
            level = min(level, 4)
            text = element.get_text(strip=True)
            if text:
                doc.add_heading(text, level=level)

        elif tag == "p":
            render_paragraph(doc, element)

        elif tag == "pre":
            code_el = element.find("code")
            text = code_el.get_text() if code_el else element.get_text()
            for line in text.splitlines():
                p = doc.add_paragraph(line)
                p.style = doc.styles["Code"]
                pf = p.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.left_indent = Cm(0.5)

        elif tag == "table":
            render_table(doc, element)

        elif tag in ("ul", "ol"):
            render_list(doc, element, ordered=(tag == "ol"))

        elif tag == "blockquote":
            text = element.get_text(strip=True)
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Cm(1.0)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        elif tag == "hr":
            pass

        elif tag in ("div", "section"):
            render_html_to_doc(doc, str(element), depth_offset)

        elif tag in ("img", "svg", "mermaid"):
            pass

        else:
            text = element.get_text(strip=True)
            if text:
                doc.add_paragraph(text)


def render_paragraph(doc, element):
    """Рендерить HTML параграф с inline-стилями."""
    p = doc.add_paragraph()
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip():
                p.add_run(text)
        elif child.name == "strong" or child.name == "b":
            r = p.add_run(child.get_text())
            r.bold = True
        elif child.name == "em" or child.name == "i":
            r = p.add_run(child.get_text())
            r.italic = True
        elif child.name == "code":
            r = p.add_run(child.get_text())
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        elif child.name == "a":
            text = child.get_text()
            if text.strip():
                r = p.add_run(text)
                r.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                r.underline = True
        elif child.name == "br":
            pass
        else:
            text = child.get_text()
            if text.strip():
                p.add_run(text)
    return p


def render_list(doc, element, ordered=False, level=0):
    """Рендерить HTML список."""
    items = element.find_all("li", recursive=False)
    for idx, li in enumerate(items):
        text = li.get_text(strip=True)
        indent = Cm(0.5 * (level + 1))
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = indent
        if ordered:
            p.style = doc.styles["List Number"] if "List Number" in [s.name for s in doc.styles] else doc.styles["Normal"]
        else:
            p.style = doc.styles["List Bullet"] if "List Bullet" in [s.name for s in doc.styles] else doc.styles["Normal"]

        sublists = li.find_all(["ul", "ol"], recursive=False)
        for sub in sublists:
            render_list(doc, sub, ordered=(sub.name == "ol"), level=level + 1)


def render_table(doc, element):
    """Рендерить HTML таблицу."""
    rows = element.find_all("tr")
    if not rows:
        return

    first_row = rows[0]
    cols = first_row.find_all(["td", "th"])
    num_cols = len(cols)
    if num_cols == 0:
        return

    table = doc.add_table(rows=min(len(rows), 100), cols=num_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        for j, cell in enumerate(cells):
            if j < num_cols:
                text = cell.get_text(strip=True)
                table.rows[i].cells[j].text = text
                if row.find("th") or cell.name == "th":
                    for p in table.rows[i].cells[j].paragraphs:
                        for run in p.runs:
                            run.bold = True
                            run.font.size = Pt(10)
                else:
                    for p in table.rows[i].cells[j].paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(10)

    doc.add_paragraph()


# ─── GRACE XML генерация ────────────────────────────────────────────

GRACE_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml"


def make_grace_manifest(modules_info):
    """Создать grace-manifest.xml."""
    parts_xml = ""
    for i, (fname, purpose, order) in enumerate([
        ("word/grace-manifest.xml", "Discovery beacon", "1"),
        ("word/grace-instructions.xml", "Agent behavioral rules", "2"),
        ("word/grace-graph.xml", "Document module map with element inventory", "3"),
        ("word/grace-contracts.xml", "Per-module and per-type editing rules", "4"),
        ("word/grace-verification.xml", "Integrity checks", "5"),
    ], 1):
        parts_xml += f'    <part-{i}><file>{fname}</file><purpose>{purpose}</purpose><read-order>{order}</read-order></part-{i}>\n'

    steps_xml = ""
    for i, (step, desc) in enumerate([
        ("Unpack the .docx", "Unpack the .docx"),
        ("Read word/grace-manifest.xml", "Read word/grace-manifest.xml"),
        ("Read word/grace-instructions.xml", "Read word/grace-instructions.xml"),
        ("Read word/grace-graph.xml", "Locate target module, check ELEMENTS for content types"),
        ("Read word/grace-contracts.xml", "Check TypeContracts for element type, then ModuleContracts for overrides"),
        ("Navigate via bookmark name or paragraph range", "Navigate via bookmark name or paragraph range"),
        ("Perform edit according to contract rules", "Perform edit according to contract rules for the specific element type"),
        ("Run verification from word/grace-verification.xml", "Run verification from word/grace-verification.xml"),
        ("Pack the .docx back", "Pack the .docx back"),
    ], 1):
        steps_xml += f'    <step-{i}>{step}</step-{i}>\n'

    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<GraceManifest VERSION="3.0.0" SCHEMA="grace-docx">
  <document-name>{DOC_NAME}</document-name>
  <document-version>{DOC_VERSION}</document-version>
  <grace-version>{GRACE_VERSION}</grace-version>
  <created>{TODAY}</created>
  <last-updated>{TODAY}</last-updated>
  <Parts>
{parts_xml}  </Parts>
  <Protocol>
{steps_xml}  </Protocol>
  <EditPolicy>
    <output-mode>new-version</output-mode>
  </EditPolicy>
  <BookmarkConvention>
    <pattern>GRACE_{{MODULE-ID}}</pattern>
    <description>Each H1 section gets a w:bookmarkStart/w:bookmarkEnd pair named GRACE_{{MODULE-ID}}.</description>
  </BookmarkConvention>
</GraceManifest>'''


def make_grace_instructions():
    """Создать grace-instructions.xml."""
    return '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<GraceInstructions VERSION="3.0.0">
  <CorePrinciples>
    <principle-1 name="contract-first">Before modifying any element, read its TypeContract in grace-contracts.xml, then check ModuleContract for overrides. Both must be satisfied.</principle-1>
    <principle-2 name="bookmark-integrity">GRACE bookmarks are navigation anchors. They must remain paired and wrap the correct section. Never delete, split, or misalign them.</principle-2>
    <principle-3 name="graph-is-current">When you add/remove/reorder content, update grace-graph.xml so future agents can navigate deterministically.</principle-3>
    <principle-4 name="verify-after-edit">After ANY edit, run the verification protocol from grace-verification.xml. If any hard-severity check fails, rollback.</principle-4>
    <principle-5 name="surgical-edits">Only change what is requested. Do not reformat, restyle, or clean up. Match existing styles exactly. Preserve all metadata attributes.</principle-5>
    <principle-6 name="element-type-awareness">Before editing a table or chart, check its type in ELEMENTS. TABLE-DATA and TABLE-STRUCT have different rules. CHART-IMAGE is readonly. CHART-NATIVE requires editing chart XML, not document.xml.</principle-6>
  </CorePrinciples>
  <EditRules>
    <rule severity="hard">Never modify w:rsidR, w:rsidRDefault, w14:paraId, w14:textId on existing elements</rule>
    <rule severity="hard">New paragraphs/runs must use same w:pStyle/w:rPr as siblings</rule>
    <rule severity="hard">Do not add/remove/reorder table columns</rule>
    <rule severity="hard">Do not promote H2 to H1 or demote H1 to H2</rule>
    <rule severity="hard">Recalculate para-range for ALL affected modules when paragraphs added/removed</rule>
    <rule severity="hard">CHART-IMAGE files in word/media/ are readonly — never modify</rule>
    <rule severity="hard">CHART-SMARTART layout-source is forbidden — only data-source text is editable</rule>
    <rule severity="hard">CHART-NATIVE data must be edited via word/charts/chartN.xml, not via document.xml drawing reference</rule>
    <rule severity="soft">Prefer append over insert</rule>
    <rule severity="soft">Batch must-sync updates in one pass</rule>
  </EditRules>
  <AntiPatterns>
    <item>Do not pretty-print document.xml — keep as single line</item>
    <item>Do not remove or rename GRACE_* bookmarks</item>
    <item>Do not delete grace-*.xml files</item>
    <item>Do not change [Content_Types].xml entries for grace-*.xml parts</item>
    <item>Do not add GRACE bookmarks without updating grace-graph.xml</item>
    <item>Do not modify content outside requested scope</item>
    <item>Do not attempt to edit CHART-IMAGE by modifying the PNG/JPG file</item>
    <item>Do not edit EMBEDDED (OLE) objects directly</item>
  </AntiPatterns>
</GraceInstructions>'''


def make_grace_graph(modules_info, total_paras, total_tables, total_h1, total_h2):
    """Создать grace-graph.xml."""
    modules_xml = ""
    for mod in modules_info:
        elements_xml = ""
        for el in mod.get("elements", []):
            attrs = " ".join(f'{k}="{v}"' for k, v in el.items())
            elements_xml += f'\n        <element {attrs}/>'
        if elements_xml:
            elements_xml = f"\n      <ELEMENTS>{elements_xml}\n      </ELEMENTS>"

        subs_xml = ""
        for sub in mod.get("subsections", []):
            subs_xml += f'\n        <sub id="{sub["id"]}">'
            subs_xml += f"\n          <heading>{sub['heading']}</heading>"
            subs_xml += f"\n          <para-start>{sub['para_start']}</para-start>"
            subs_xml += f"\n          <para-end>{sub['para_end']}</para-end>"
            subs_xml += "\n        </sub>"
        if subs_xml:
            subs_xml = f"\n      <SubSections>{subs_xml}\n      </SubSections>"

        mod_id = mod["id"]
        modules_xml += f'''
    <{mod_id}>
      <n>{mod["heading"]}</n>
      <TYPE>{mod["type"]}</TYPE>
      <BOOKMARK>GRACE_{mod_id}</BOOKMARK>
      <PARA-START>{mod["para_start"]}</PARA-START>
      <PARA-END>{mod["para_end"]}</PARA-END>{subs_xml}{elements_xml}
    </{mod_id}>'''

    cross_xml = ""
    for link in get_cross_links(modules_info):
        cross_xml += f'''
    <link>
      <from>{link["from"]}</from>
      <to>{link["to"]}</to>
      <relation>{link["relation"]}</relation>
    </link>'''

    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<GraceGraph VERSION="3.0.0">
  <DocumentMeta>
    <total-paragraphs>{total_paras}</total-paragraphs>
    <total-tables>{total_tables}</total-tables>
    <total-h1>{total_h1}</total-h1>
    <total-h2>{total_h2}</total-h2>
  </DocumentMeta>
  <Modules>{modules_xml}
  </Modules>
  <CrossLinks>{cross_xml}
  </CrossLinks>
</GraceGraph>'''


def make_grace_contracts(modules_info):
    """Создать grace-contracts.xml."""
    module_contracts = ""
    for mod in modules_info:
        parent_type = "C-NARRATIVE" if mod["type"] in ("NARRATIVE", "META", "NAVIGATION", "REFERENCE") else \
                     "C-TABLE-DATA" if mod["type"] == "DATA" else "C-MIXED"
        module_contracts += f'''
    <C-{mod["id"]} inherits="{parent_type}">
      <description>{mod["heading"]} — {mod["type"]}-type section</description>
      <can-edit>
        <item>Add paragraphs after existing content, modify text runs</item>
      </can-edit>
      <cannot-edit>
        <item>Change heading styles or structure</item>
      </cannot-edit>
    </C-{mod["id"]}>'''

    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<GraceContracts VERSION="3.0.0">
  <GlobalRules>
    <rule severity="hard">Never remove or merge GRACE bookmark pairs</rule>
    <rule severity="hard">Table column structure is immutable — do not add, remove, or reorder columns</rule>
    <rule severity="hard">Never change w:rsidR, w14:paraId on existing paragraphs</rule>
    <rule severity="hard">CHART-IMAGE and VISUAL-IMAGE are readonly</rule>
    <rule severity="hard">CHART-SMARTART topology is immutable</rule>
    <rule severity="hard">EMBEDDED objects are readonly</rule>
    <rule severity="soft">Prefer adding new paragraphs over modifying existing ones</rule>
    <rule severity="soft">Match surrounding paragraph style when adding content</rule>
  </GlobalRules>
  <TypeContracts>
    <C-NARRATIVE>
      <description>Prose sections: paragraphs, bullets, numbered lists</description>
      <can-edit>Add paragraphs after existing content, modify text runs, update numbered/bulleted items</can-edit>
      <cannot-edit>Change heading styles, modify list numbering format, alter paragraph indentation</cannot-edit>
    </C-NARRATIVE>
    <C-TABLE-DATA>
      <description>Tables with a header row and data rows below</description>
      <can-edit>Add rows at the end, update cell values in data rows</can-edit>
      <cannot-edit>Modify header row text or formatting, add or remove columns, merge cells</cannot-edit>
    </C-TABLE-DATA>
    <C-TABLE-STRUCT>
      <description>Structural tables: RACI matrices, comparison grids</description>
      <can-edit>Update cell text values within existing structure</can-edit>
      <cannot-edit>Add or remove rows or columns, change cell merging</cannot-edit>
    </C-TABLE-STRUCT>
    <C-MIXED>
      <description>Mixed prose and data content</description>
      <can-edit>Add paragraphs, update table data rows, modify prose text</can-edit>
      <cannot-edit>Change table headers, modify structure</cannot-edit>
    </C-MIXED>
  </TypeContracts>
  <ModuleContracts>{module_contracts}
  </ModuleContracts>
</GraceContracts>'''


def make_grace_verification(num_modules):
    """Создать grace-verification.xml."""
    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<GraceVerification VERSION="3.0.0">
  <StructuralInvariants>
    <invariant id="bookmark-balance" severity="hard">
      Count bookmarkStart with name starting "GRACE_". Count bookmarkEnd. Must be equal. Expected: {num_modules} pairs.
    </invariant>
    <invariant id="heading-hierarchy" severity="hard">
      For each H2, a preceding H1 must exist. H1 sections must not nest.
    </invariant>
    <invariant id="grace-xml-valid" severity="hard">
      All grace-*.xml files must parse as well-formed XML without error.
    </invariant>
    <invariant id="graph-covers-all-h1" severity="hard">
      Every H1 heading in document.xml must have a matching module in grace-graph.xml. Expected: {num_modules}.
    </invariant>
    <invariant id="table-column-consistency" severity="hard">
      For each w:tbl, count w:tc in each w:tr. All rows must match.
    </invariant>
  </StructuralInvariants>
  <PostEditChecks>
    <check id="paragraph-range-accuracy">After adding/removing paragraphs, re-scan heading positions and update all para-range values in grace-graph.xml.</check>
    <check id="bookmark-intact">After edits near a bookmark boundary, verify bookmarkStart/bookmarkEnd still wrap the expected H1 heading.</check>
    <check id="styles-preserved">When modifying text, compare w:pPr and w:rPr before/after. Only w:t should change.</check>
    <check id="elements-inventory-current">After adding or removing tables, update the ELEMENTS block for the affected module.</check>
  </PostEditChecks>
  <ValidationProtocol>
    <step>Run all StructuralInvariants before edit</step>
    <step>If any hard-severity fails — STOP, do not proceed</step>
    <step>Perform edit according to TypeContract for the element type</step>
    <step>Run all StructuralInvariants again</step>
    <step>Run relevant PostEditChecks</step>
    <step>If any hard check fails — ROLLBACK</step>
    <step>Update grace-graph.xml if structure changed</step>
    <step>Pack document</step>
  </ValidationProtocol>
</GraceVerification>'''


def get_cross_links(modules_info):
    """Определить перекрёстные ссылки между модулями."""
    links = []
    parent_map = {}
    for mod in modules_info:
        if mod.get("parent"):
            links.append({
                "from": mod["id"],
                "to": mod["parent"],
                "relation": f"references: подраздел {mod['heading']} относится к {mod['parent_heading']}"
            })
    return links


# ─── Основной скрипт генерации ──────────────────────────────────────

def collect_chapters(chapters_dict, depth=0, parent_id=None, parent_heading=None):
    """Собрать плоский список глав для генерации."""
    result = []
    for heading, source in chapters_dict.items():
        if isinstance(source, str):
            result.append({
                "heading": heading,
                "source": source,
                "depth": depth,
                "parent_id": parent_id,
                "parent_heading": parent_heading,
            })
        elif isinstance(source, dict):
            result.append({
                "heading": heading,
                "source": None,
                "depth": depth,
                "parent_id": parent_id,
                "parent_heading": parent_heading,
                "is_container": True,
            })
            result.extend(collect_chapters(source, depth + 1, parent_id=heading, parent_heading=heading))
    return result


def main():
    print("=== Генерация Word-документа с GRACE-разметкой ===\n")

    doc = Document()
    setup_styles(doc)

    # ─── Титульная страница ──────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("СТАНДАРТЫ РАЗРАБОТКИ")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle_p.add_run("Гид по внутренним процессам и стандартам\nпроектной команды")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()

    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ver_p.add_run(f"Версия {DOC_VERSION}")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = date_p.add_run(TODAY)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()
    doc.add_paragraph()

    org_p = doc.add_paragraph()
    org_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = org_p.add_run("yellow-hammer / dev-rules")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ─── Оглавление (заполнитель) ────────────────────────
    toc_heading = doc.add_heading("Оглавление", level=1)

    chapters = collect_chapters(CHAPTERS)
    existing_ids = set()
    modules_info = []
    all_table_count = 0
    toc_entries = []

    for idx, ch in enumerate(chapters):
        mid = derive_module_id(ch["heading"], existing_ids)
        existing_ids.add(mid)
        ch["module_id"] = mid
        toc_entries.append((ch["heading"], ch["depth"], mid))

        indent = "    " * ch["depth"]
        toc_p = doc.add_paragraph()
        toc_p.paragraph_format.left_indent = Cm(ch["depth"] * 0.75)
        r = toc_p.add_run(f"{ch['heading']}")
        if ch["depth"] == 0:
            r.bold = True
            r.font.size = Pt(12)
        else:
            r.font.size = Pt(11)

    doc.add_page_break()

    # ─── Генерация глав ──────────────────────────────────
    para_counter = 0
    table_counter = 0
    h1_counter = 0
    h2_counter = 0
    bookmark_id = 100

    for ch in chapters:
        mid = ch["module_id"]
        depth = ch["depth"]
        heading_level = depth + 1

        # Вставить закладку GRACE перед заголовком
        heading_para = doc.add_heading(ch["heading"], level=heading_level)

        # Вставить bookmarkStart
        bm_start = OxmlElement("w:bookmarkStart")
        bm_start.set(qn("w:id"), str(bookmark_id))
        bm_start.set(qn("w:name"), f"GRACE_{mid}")
        heading_para._element.insert(0, bm_start)

        # Bookmark placeholder — будет закрыт перед следующим модулем
        ch["bookmark_start_id"] = bookmark_id
        ch["para_start"] = para_counter
        bookmark_id += 1
        para_counter += 1

        if heading_level == 1:
            h1_counter += 1
        elif heading_level == 2:
            h2_counter += 1

        mod_type = "NARRATIVE"
        mod_elements = []

        if ch.get("source"):
            title, md_text = read_md(ch["source"])
            if md_text:
                html = md_to_html(md_text)
                mod_type = classify_module_type(html)

                # Подсчитать таблицы в HTML
                html_tables = html.lower().count("<table>")
                table_counter += html_tables
                for t_idx in range(html_tables):
                    mod_elements.append({
                        "type": "TABLE-DATA",
                        "para-index": str(para_counter + t_idx),
                        "columns": "0",
                        "rows": "0",
                    })

                # Рендерить контент
                render_html_to_doc(doc, html, depth_offset=depth)
                para_counter += len(html.split("<p>")) + len(html.split("<h")) + len(html.split("<li>"))

        else:
            mod_type = "NAVIGATION"

        ch["type"] = mod_type
        ch["elements"] = mod_elements
        ch["para_end"] = para_counter - 1

        # Добавить bookmarkEnd в конец секции
        bm_end = OxmlElement("w:bookmarkEnd")
        bm_end.set(qn("w:id"), str(ch["bookmark_start_id"]))
        last_para = doc.paragraphs[-1]._element
        last_para.addnext(bm_end)

        modules_info.append({
            "id": mid,
            "heading": ch["heading"],
            "type": mod_type,
            "para_start": ch["para_start"],
            "para_end": ch["para_end"],
            "elements": mod_elements,
            "subsections": [],
            "parent": ch.get("parent_id"),
            "parent_heading": ch.get("parent_heading"),
        })

    # ─── Сохранить документ ──────────────────────────────
    doc.save(str(OUTPUT_PATH))
    print(f"Документ сохранён: {OUTPUT_PATH}")

    # ─── Инъекция GRACE XML-частей ───────────────────────
    print("Инъекция GRACE XML-частей...")

    import zipfile
    import shutil
    import tempfile

    tmp_dir = Path(tempfile.mkdtemp())
    grace_dir = tmp_dir / "word"
    grace_dir.mkdir(exist_ok=True)

    # Распаковать docx
    with zipfile.ZipFile(str(OUTPUT_PATH), "r") as z:
        z.extractall(str(tmp_dir))

    # Записать GRACE XML-файлы
    (grace_dir / "grace-manifest.xml").write_text(
        make_grace_manifest(modules_info), encoding="utf-8"
    )
    (grace_dir / "grace-instructions.xml").write_text(
        make_grace_instructions(), encoding="utf-8"
    )
    (grace_dir / "grace-graph.xml").write_text(
        make_grace_graph(modules_info, para_counter, table_counter, h1_counter, h2_counter),
        encoding="utf-8",
    )
    (grace_dir / "grace-contracts.xml").write_text(
        make_grace_contracts(modules_info), encoding="utf-8"
    )
    (grace_dir / "grace-verification.xml").write_text(
        make_grace_verification(len(modules_info)), encoding="utf-8"
    )

    # Обновить [Content_Types].xml
    ct_path = tmp_dir / "[Content_Types].xml"
    ct_content = ct_path.read_text(encoding="utf-8")
    grace_overrides = ""
    for name in ["grace-manifest.xml", "grace-instructions.xml", "grace-graph.xml",
                 "grace-contracts.xml", "grace-verification.xml"]:
        grace_overrides += f'<Override PartName="/word/{name}" ContentType="application/xml"/>\n'
    ct_content = ct_content.replace("</Types>", f"{grace_overrides}</Types>")
    ct_path.write_text(ct_content, encoding="utf-8")

    # Обновить word/_rels/document.xml.rels
    rels_path = tmp_dir / "word" / "_rels" / "document.xml.rels"
    if rels_path.exists():
        rels_content = rels_path.read_text(encoding="utf-8")
        grace_rels = ""
        for i, name in enumerate(["grace-manifest.xml", "grace-instructions.xml",
                                   "grace-graph.xml", "grace-contracts.xml",
                                   "grace-verification.xml"], 1):
            rid = f"rIdGrace{i}"
            if rid not in rels_content:
                grace_rels += f'<Relationship Id="{rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/customXml" Target="{name}"/>\n'
        rels_content = rels_content.replace("</Relationships>", f"{grace_rels}</Relationships>")
        rels_path.write_text(rels_content, encoding="utf-8")

    # Упаковать обратно
    grace_output = OUTPUT_PATH.parent / f"GRACE_{OUTPUT_PATH.name}"
    with zipfile.ZipFile(str(grace_output), "w", zipfile.ZIP_DEFLATED) as zout:
        for file_path in sorted(tmp_dir.rglob("*")):
            if file_path.is_file():
                arcname = file_path.relative_to(tmp_dir)
                zout.write(str(file_path), str(arcname))

    shutil.rmtree(str(tmp_dir))

    import sys as _sys, io as _io
    _sys.stdout = _io.TextIOWrapper(_sys.stdout.buffer, encoding="utf-8", errors="replace")

    print(f"""
=============================================
GRACE-DOCX Bootstrap Complete  [v3]
=============================================
Document:        {DOC_NAME}
Version:         {DOC_VERSION}
Modules:         {len(modules_info)} identified, {len(modules_info)} bookmarked
XML parts:       5 injected
Bookmarks:       {len(modules_info)} pairs injected
---------------------------------------------
Module IDs:""")

    for mod in modules_info:
        print(f"  {mod['id']:12s}  {mod['type']:12s}  {mod['heading']}")

    print(f"""---------------------------------------------
Output files:
  Base:    {OUTPUT_PATH}
  GRACE:   {grace_output}
=============================================
""")


if __name__ == "__main__":
    main()
